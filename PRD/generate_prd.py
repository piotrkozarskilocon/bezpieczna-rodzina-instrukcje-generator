"""
Generuje dokument PRD (Word) dla integracji generatora instrukcji z Claude API.
Uruchomienie:  python PRD/generate_prd.py

Plik wyjściowy: PRD/PRD - Integracja Claude API.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import date


OUTPUT_PATH = Path(__file__).parent / "PRD - Integracja Claude API.docx"


def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style_normal(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_meta_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.width = Cm(4.5)
        c1.width = Cm(12)
        c0.text = k
        c1.text = v
        for run in c0.paragraphs[0].runs:
            run.bold = True
        set_cell_bg(c0, "F2F4F7")


def add_table_with_header(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
        set_cell_bg(cell, "1F2937")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.cell(ri + 1, ci).text = val

    if col_widths_cm:
        for row in table.rows:
            for cell, w in zip(row.cells, col_widths_cm):
                cell.width = Cm(w)


def H1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)


def H2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def H3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)


def P(doc, text):
    doc.add_paragraph(text)


def BULLETS(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def NUMBERED(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def CODE(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p_pr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), "CBD5E1")
        pBdr.append(b)
    p_pr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F8FAFC")
    p_pr.append(shd)


# ─────────────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    style_normal(doc)

    # ─── Strona tytułowa ────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("PRD — Integracja Claude API")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Generator instrukcji obsługi v4 (AI-first)")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    add_horizontal_line(subtitle)

    add_meta_table(doc, [
        ("Produkt", "Generator instrukcji obsługi Locon / Bezpieczna Rodzina"),
        ("Wersja PRD", "1.0"),
        ("Data", date.today().isoformat()),
        ("Autor", "Piotr Kozarski (PKZ / Locon Sp. z o.o.)"),
        ("Status", "Draft do akceptacji"),
        ("Repo", "github.com/piotrkozarskilocon/bezpieczna-rodzina-instrukcje-generator"),
        ("Środowisko", "Vercel (Next.js 16) + Supabase Postgres + Anthropic API"),
    ])

    doc.add_paragraph()

    # ─── 1. Streszczenie zarządcze ──────────────────────────────────────────
    H1(doc, "1. Streszczenie zarządcze")
    P(doc,
      "Generator instrukcji obsługi v4 jest narzędziem wewnętrznym do tworzenia "
      "wielojęzycznych dokumentów drukowanych dla urządzeń Locon (smartwatche, "
      "opaski seniorskie, trackery zwierzęce). W obecnym stanie korzysta z modelu "
      "Claude w trybie manualnym — użytkownik kopiuje wygenerowane prompty, wkleja "
      "je w okno claude.ai, a zwrócone artefakty JSON wkleja z powrotem do aplikacji. "
      "Niniejszy PRD opisuje przejście na pełną integrację z Anthropic API, dzięki "
      "której każdy z pięciu obecnych workflow-ów stanie się jednokliknięciowy.")
    P(doc,
      "Po wdrożeniu integracji koszt operacyjny przeniesie się z czasu autora "
      "instrukcji (~30–45 min ręcznego kopiowania per nowy projekt + każda iteracja) "
      "na koszt API (~2–4 USD per pełny projekt + ~0,05–0,20 USD per edycja strony). "
      "Time-to-first-draft skróci się z ~1 godziny do ~2 minut.")

    # ─── 2. Tło i problem ───────────────────────────────────────────────────
    H1(doc, "2. Tło i problem")

    H2(doc, "2.1 Stan obecny")
    P(doc,
      "Generator obsługuje pełen cykl życia dokumentu: konfiguracja w wizardzie "
      "(typ dokumentu, typ urządzenia, funkcje), generacja struktury stron + "
      "elementów, edycja w canvas-owym edytorze (drag/drop/properties), tłumaczenia "
      "do 6 języków, eksport do wektorowego PDF. Wszystko trzymane w Supabase "
      "(gen4_projects, gen4_pages, gen4_elements, gen4_translations).")
    P(doc, "Pięć workflow-ów wymaga dziś manualnego copy/paste do claude.ai:")
    add_table_with_header(doc,
        headers=["#", "Workflow", "Endpoint prompt", "Endpoint import"],
        rows=[
            ["1", "Pełna generacja projektu", "GET /api/v4/projects/[id]/prompt", "POST /api/v4/projects/[id]/import"],
            ["2", "Edycja pojedynczej strony (Assistant AI)", "POST /api/v4/pages/[pageId]/edit-prompt", "POST /api/v4/pages/[pageId]/replace-elements"],
            ["3", "Zastosowanie design systemu do projektu", "POST /api/v4/projects/[id]/apply-design-prompt", "POST /api/v4/projects/[id]/import"],
            ["4", "Zastosowanie design systemu do strony", "POST /api/v4/pages/[pageId]/apply-design-prompt", "POST /api/v4/pages/[pageId]/replace-elements"],
            ["5", "Tłumaczenie projektu na inny język", "GET /api/v4/projects/[id]/translate-prompt?lang=bg", "POST /api/v4/projects/[id]/translations"],
        ],
        col_widths_cm=[0.8, 5.5, 6, 5])

    H2(doc, "2.2 Pain points trybu manualnego")
    BULLETS(doc, [
        "Czas: jeden cykl prompt → claude.ai → JSON → import zajmuje 3–10 minut, "
        "a pełne wygenerowanie + 4–6 iteracji per projekt to ~30–60 min wyłącznie "
        "na kopiowaniu, nie na właściwej pracy nad treścią.",
        "Błędy klejenia: użytkownik łatwo wkleja niekompletny JSON (urwany fence), "
        "przez co import pada — kosztuje to dodatkowe minuty na ponowne uruchomienie "
        "rozmowy w claude.ai.",
        "Brak wersjonowania promptów: każdy nowy chat w claude.ai oznacza inny "
        "kontekst, brak deterministyczności. Trudno odtworzyć identyczny wynik "
        "tygodnie później.",
        "Brak telemetrii: nie wiemy ile tokenów / pieniędzy idzie na jeden projekt, "
        "nie umiemy mierzyć regresji jakości promptu.",
        "Workflow nie skaluje się: 17 modeli urządzeń × 6 języków = 102 projekty. "
        "Manualny tryb obsłuży 1-2 dziennie. Auto-tryb obsłuży tyle, ile wystawi "
        "rate limit Anthropic API (3,000 RPM / 400k TPM dla tier 4)."
    ])

    H2(doc, "2.3 Dlaczego teraz")
    BULLETS(doc, [
        "Cała infrastruktura aplikacji (Next.js 16 + Supabase + Vercel) jest "
        "gotowa: endpoint generate już ma fallback do API (callClaude w lib/anthropic.ts), "
        "wystarczy ANTHROPIC_API_KEY w env i przepiąć resztę endpointów.",
        "Locon planuje wydanie 3 nowych modeli w 2026 (Slay AI, Goat 2, opaska "
        "senior Premium) — każdy wymaga komplecie PL+EN+BG+RO+MK+SQ+HR. To 21 "
        "dokumentów które bez API zatłoczą backlog.",
        "Cena tokenów Claude Sonnet 4.6 (3 USD/M input, 15 USD/M output) jest "
        "dziś niższa niż jeden kawowy wieczór czasu pracownika; ROI > 100x "
        "z perspektywy czasu autora dokumentu."
    ])

    # ─── 3. Cele i metryki sukcesu ──────────────────────────────────────────
    H1(doc, "3. Cele i metryki sukcesu")

    H2(doc, "3.1 Cele biznesowe")
    BULLETS(doc, [
        "Skrócić czas produkcji pierwszej wersji instrukcji z 4–8 godzin do < 30 minut.",
        "Umożliwić batch-generację (wszystkie 6 języków, wszystkie typy dokumentów "
        "dla nowego modelu) w jednym sprincie zamiast w 2 tygodnie pracy.",
        "Zbudować spójną telemetrię kosztów: koszt API per projekt, per użytkownik, "
        "per typ dokumentu — dla decyzji wewnętrznych (czy używać Sonnet czy Haiku).",
    ])

    H2(doc, "3.2 Wskaźniki sukcesu (KPI)")
    add_table_with_header(doc,
        headers=["KPI", "Stan obecny", "Cel po wdrożeniu", "Sposób pomiaru"],
        rows=[
            ["Time-to-first-draft", "≈ 60 min", "≤ 3 min", "Telemetria w gen4_ai_history.latency_ms"],
            ["% prób importu z błędem JSON", "≈ 10–15%", "< 1%", "Logi import endpointu (status 400)"],
            ["Koszt operacyjny / projekt", "≈ 90 min pracy autora", "≤ 5 USD API + 10 min pracy autora", "Sumaryczne ai_log + czas w wizardzie"],
            ["Liczba projektów / tydzień", "≈ 2–3", "≥ 15", "COUNT(*) gen4_projects.created_at"],
            ["% projektów z pełnym tłumaczeniem 6 lang", "≈ 20%", "≥ 80%", "Coverage gen4_translations per project"],
        ],
        col_widths_cm=[4.5, 3.5, 4.5, 4.5])

    # ─── 4. Persony i scenariusze ──────────────────────────────────────────
    H1(doc, "4. Persony i scenariusze użycia")
    H2(doc, "4.1 Persona — Product Manager (autor instrukcji)")
    P(doc,
      "Piotr / inny PM w Locon. Wie czego potrzebuje w treści (kroki uruchomienia, "
      "info SAR, warunki gwarancji), ale nie chce ręcznie układać layoutu ani "
      "kopiować promptów. Korzysta z generatora ~5x w tygodniu przy nowym modelu, "
      "potem w trybie maintenance kilka razy w miesiącu (poprawki gwarancji, "
      "aktualizacja kontaktu, nowy język).")

    H2(doc, "4.2 Główne scenariusze")
    NUMBERED(doc, [
        "Tworzenie nowego dokumentu od zera: wybór (typ dokumentu × typ urządzenia × "
        "model) w wizardzie → \"Generuj\" → po ~30 s wyświetla się gotowy projekt "
        "ze wszystkimi wymaganymi sekcjami i widocznymi placeholderami ⚠️ DO "
        "UZUPEŁNIENIA.",
        "Iteracyjna edycja strony przez Assistant AI: \"zrób cover bardziej elegancki\" "
        "→ jeden klik → po ~10 s strona jest przepisana.",
        "Zastosowanie design systemu Bezpieczna Rodzina do projektu opartego o "
        "design Stop Hejt: jeden klik \"Zastosuj DS do całego projektu\" → po "
        "~60 s wszystkie strony używają tokens z DS.",
        "Tłumaczenie projektu na bułgarski/macedoński/inny: wybór języka w panelu "
        "tłumaczeń → jeden klik → po ~20 s pojawiają się rzędy w gen4_translations, "
        "PDF Export może je użyć.",
        "Powrót do manualu (\"fallback\"): klucz API wyczerpany / Anthropic ma "
        "outage → przycisk \"Kopiuj prompt\" wciąż działa, użytkownik może "
        "ratunkowo wkleić w claude.ai."
    ])

    # ─── 5. Zakres ─────────────────────────────────────────────────────────
    H1(doc, "5. Zakres")
    H2(doc, "5.1 In scope (must have)")
    BULLETS(doc, [
        "ANTHROPIC_API_KEY skonfigurowany w env Vercel (Production + Preview).",
        "Przepięcie wszystkich 5 workflow-ów (generacja projektu, edycja strony, "
        "apply DS projekt, apply DS strona, tłumaczenia) na auto-tryb z API.",
        "Zachowanie manualnego trybu jako fallback (UI pokazuje przycisk \"Kopiuj "
        "prompt\" obok przycisku \"Generuj z AI\").",
        "Telemetria: każde wywołanie zapisuje się w gen4_ai_history (model, "
        "input_tokens, output_tokens, latency_ms, koszt USD).",
        "Obsługa błędów: 429 (rate limit) → retry z exponential backoff, 5xx → "
        "retry 1×, inne → komunikat do użytkownika + zachowanie promptu by mógł "
        "ręcznie skopiować.",
        "Selekcja modelu per workflow: Sonnet 4.6 dla pełnej generacji + apply DS, "
        "Haiku 4.5 dla edycji pojedynczej strony i tłumaczeń.",
        "Dashboard kosztowy w sekcji \"Project info\" — pokazuje sumę USD wydaną "
        "na dany projekt.",
        "Rate limiting po stronie aplikacji: max 1 wywołanie / projekt / 5 s "
        "(zapobiega przypadkowemu double-click)."
    ])

    H2(doc, "5.2 Out of scope (faza 1)")
    BULLETS(doc, [
        "Per-user API key (zostajemy z jednym kluczem app-level — wystarczy "
        "dla 1-3 użytkowników wewnętrznych).",
        "Streaming responses w UI (na razie blokujemy UI na loaderze; streaming "
        "ma sens dopiero gdy generacja > 30 s).",
        "Generowanie obrazów (Claude nie generuje, a Imagen/DALL-E to osobny "
        "projekt — placeholdery zostają).",
        "Prompt caching API (Anthropic) — wprowadzić w fazie 2, gdy widzimy "
        "powtarzające się sekcje promptów (np. design system).",
        "Multi-tenant SaaS — produkt jest narzędziem wewnętrznym Locon."
    ])

    # ─── 6. Wymagania funkcjonalne ─────────────────────────────────────────
    H1(doc, "6. Wymagania funkcjonalne")

    H2(doc, "6.1 FR-1: Wywołanie API z poziomu wizardu")
    P(doc, "Gdy użytkownik klika \"Stwórz projekt + wygeneruj z AI\":")
    NUMBERED(doc, [
        "Backend insert do gen4_projects (status = 'generating').",
        "Wywołanie callClaude({ system, user, model: INITIAL_MODEL, maxTokens: 32000 }).",
        "Parsowanie odpowiedzi (parseJsonFromAi + validateGenerated).",
        "bulkInsertGeneratedProject — strony + elementy + tytuły.",
        "Update gen4_projects.status = 'ready', append wpisu do ai_log.",
        "Redirect na /ai/projects/[id] (z 'ready' state, edytor otwarty)."
    ])
    P(doc, "Akceptanca: cykl < 60 s przy projekcie ~10 stron, brak błędu 400 import.")

    H2(doc, "6.2 FR-2: Auto-tryb dla Assistant AI (edycja strony)")
    P(doc, "Nowy endpoint POST /api/v4/pages/[pageId]/ai-edit przyjmuje { instruction }:")
    NUMBERED(doc, [
        "Buduje prompt jak dziś (buildPageEditPrompt).",
        "Wywołuje callClaude({ system, user, model: EDIT_MODEL, maxTokens: 16000 }).",
        "Parsuje, waliduje, wykonuje replacePageElements wewnętrznie.",
        "Zwraca { elements: <count>, cost_usd } bez konieczności drugiego requestu."
    ])
    P(doc, "UI: przycisk \"✨ Zastosuj\" obok obecnego \"Kopiuj prompt\". Domyślnie "
           "aktywny tryb auto, gdy klucz API jest skonfigurowany. Akceptanca: jedno "
           "kliknięcie → strona zaktualizowana w < 15 s.")

    H2(doc, "6.3 FR-3: Auto-tryb dla Apply Design System")
    P(doc, "Analogicznie dla obu poziomów (projekt i strona). Nowe endpointy:")
    BULLETS(doc, [
        "POST /api/v4/projects/[id]/apply-design (body: { ds_id, instruction? }) "
        "→ Sonnet 4.6, max_tokens 32000",
        "POST /api/v4/pages/[pageId]/apply-design (body: { ds_id, instruction? }) "
        "→ Haiku 4.5, max_tokens 16000",
    ])
    P(doc, "Akceptanca: projekt-wide < 90 s, strona < 15 s.")

    H2(doc, "6.4 FR-4: Auto-tryb tłumaczeń")
    P(doc, "Nowy endpoint POST /api/v4/projects/[id]/translate?lang=<lang>:")
    NUMBERED(doc, [
        "Buduje prompt (buildTranslationPrompt).",
        "callClaude({ ..., model: EDIT_MODEL }).",
        "Parsuje (parseTranslationResponse).",
        "Upsert do gen4_translations.",
        "Zwraca { translations: <count>, cost_usd }."
    ])
    P(doc, "Akceptanca: pełne tłumaczenie projektu 50–100 elementów < 30 s.")

    H2(doc, "6.5 FR-5: Dual-mode UI (auto + manual)")
    P(doc, "Każdy z 5 workflow-ów ma w UI dwa przyciski obok siebie:")
    BULLETS(doc, [
        "[✨ Wygeneruj z AI (auto)] — wywołuje endpoint API, blokuje UI loader-em.",
        "[📋 Kopiuj prompt (manual)] — jak dziś, otwiera tekst do skopiowania."
    ])
    P(doc, "Akceptanca: jeśli klucz API nie jest skonfigurowany lub wywołanie API "
           "padło, przycisk auto pokazuje błąd ale przycisk manual pozostaje "
           "funkcjonalny.")

    H2(doc, "6.6 FR-6: Dashboard kosztowy projektu")
    P(doc, "Sekcja \"Project info\" pokazuje:")
    BULLETS(doc, [
        "Sumaryczny koszt USD (z gen4_ai_history za wszystkie wywołania na projekcie).",
        "Liczbę wywołań per typ (generation, edit, apply_ds, translate).",
        "Liczbę tokenów input/output (sumarycznie).",
        "Ostatnie wywołanie: model, latency, koszt."
    ])

    H2(doc, "6.7 FR-7: Telemetria w gen4_ai_history")
    P(doc, "Każde wywołanie API generuje wiersz w gen4_ai_history z polami:")
    add_table_with_header(doc,
        headers=["Pole", "Typ", "Opis"],
        rows=[
            ["project_id", "uuid", "FK do gen4_projects"],
            ["role", "text", "'assistant' dla odpowiedzi AI"],
            ["content", "text", "skrót odpowiedzi (pierwsze 500 znaków)"],
            ["structured", "jsonb", "metadane wywołania: workflow_type, page_id, ds_id, lang, success"],
            ["model", "text", "claude-sonnet-4-6 / claude-haiku-4-5-...?"],
            ["input_tokens", "int", "liczba tokenów input"],
            ["output_tokens", "int", "liczba tokenów output"],
            ["latency_ms", "int", "czas wywołania API"],
            ["created_at", "timestamptz", "moment wywołania"],
        ],
        col_widths_cm=[3.5, 2.5, 11])
    P(doc, "Koszt USD obliczany on-the-fly (cost_per_input_token * input_tokens + "
           "cost_per_output_token * output_tokens) — tabela cen w lib/anthropic.ts.")

    # ─── 7. Wymagania niefunkcjonalne ──────────────────────────────────────
    H1(doc, "7. Wymagania niefunkcjonalne")

    H2(doc, "7.1 Bezpieczeństwo")
    BULLETS(doc, [
        "ANTHROPIC_API_KEY zapisany wyłącznie jako Encrypted env var w Vercel "
        "(Production + Preview). NIE w .env.example, NIE w repo.",
        "Klucz nie jest eksponowany do klienta — wszystkie wywołania API idą "
        "przez Next.js route handlers (server-only).",
        "Authentication: każdy z endpointów AI sprawdza JWT cookie z hub site "
        "(authenticate() z lib/auth.ts). Brak public endpointu.",
        "Owner check: każdy endpoint weryfikuje że auth.email == owner_email "
        "projektu (zapobiega cross-tenant access nawet wewnątrz Locon).",
        "Rate limiting na poziomie API: per user email max 60 wywołań / godzinę. "
        "Implementacja w Supabase z LIMIT i 60-min sliding window na "
        "gen4_ai_history."
    ])

    H2(doc, "7.2 Wydajność")
    add_table_with_header(doc,
        headers=["Workflow", "Model", "Max tokens", "Cel latency p50", "Cel latency p95"],
        rows=[
            ["Pełna generacja projektu (~10 stron)", "Sonnet 4.6", "32000", "30 s", "60 s"],
            ["Edycja pojedynczej strony", "Haiku 4.5", "16000", "8 s", "20 s"],
            ["Apply DS — projekt", "Sonnet 4.6", "32000", "45 s", "90 s"],
            ["Apply DS — strona", "Haiku 4.5", "16000", "8 s", "20 s"],
            ["Tłumaczenie projektu", "Haiku 4.5", "32000", "15 s", "40 s"],
        ],
        col_widths_cm=[5, 3, 2, 2.5, 2.5])

    H2(doc, "7.3 Koszty (szacunek)")
    P(doc, "Sonnet 4.6: 3 USD / 1M input tokens, 15 USD / 1M output tokens.")
    P(doc, "Haiku 4.5: 1 USD / 1M input tokens, 5 USD / 1M output tokens.")
    add_table_with_header(doc,
        headers=["Wywołanie", "~ input tokens", "~ output tokens", "Koszt USD"],
        rows=[
            ["Pełna generacja (Sonnet 4.6)", "8 000", "12 000", "0,20"],
            ["Edycja strony (Haiku 4.5)", "3 000", "4 000", "0,02"],
            ["Apply DS projekt (Sonnet 4.6)", "15 000", "18 000", "0,32"],
            ["Apply DS strona (Haiku 4.5)", "4 000", "5 000", "0,03"],
            ["Tłumaczenie projektu (Haiku 4.5)", "10 000", "12 000", "0,07"],
        ],
        col_widths_cm=[5.5, 3, 3, 3])
    P(doc, "Szacunkowy koszt pełnego projektu (generacja + 4 edycje + 1 apply DS + "
           "6 tłumaczeń): ~ 0,85 USD. Budżet miesięczny przy 30 nowych projektach: "
           "~25 USD. Budżet roczny: ~300 USD.")

    H2(doc, "7.4 Niezawodność")
    BULLETS(doc, [
        "Retry strategy: 429 → exponential backoff (1s, 2s, 4s, max 3 próby). "
        "5xx → retry 1× po 2s. Inne błędy → bez retry, prezentacja błędu.",
        "Timeout: 120 s dla pełnej generacji (Vercel maxDuration), 30 s dla "
        "operacji edycyjnych.",
        "Fallback to manual: jeśli wszystkie próby padły, użytkownik widzi "
        "komunikat \"Generacja AI niedostępna — użyj trybu manualnego\" i "
        "przycisk \"📋 Kopiuj prompt\".",
        "Status reporting: w gen4_projects.status = 'error' z opisem błędu w "
        "gen4_ai_history.structured.error_message."
    ])

    H2(doc, "7.5 Monitorowanie")
    BULLETS(doc, [
        "Logi Vercel Functions: każde wywołanie API loguje workflow_type + "
        "project_id + duration.",
        "Dashboard SQL (Supabase Reports): tygodniowy widok kosztów, p95 "
        "latency, % błędów per workflow.",
        "Alert: jeśli koszt dzienny > 10 USD, wysłać e-mail do admina (Edge "
        "Function + cron). Faza 1 może użyć ręcznego check raz w tygodniu."
    ])

    # ─── 8. Architektura ───────────────────────────────────────────────────
    H1(doc, "8. Architektura integracji")

    H2(doc, "8.1 Komponenty")
    BULLETS(doc, [
        "lib/anthropic.ts — istniejący klient @anthropic-ai/sdk + callClaude(). "
        "Wymaga rozszerzenia o retry logic i obliczanie kosztu.",
        "lib/v4Generate.ts / v4Edit.ts / v4ApplyDs.ts / v4Translate.ts — gotowe "
        "buildery promptów. Bez zmian funkcjonalnych.",
        "app/api/v4/* — istniejące endpointy zostają (manual). Dodajemy nowe "
        "endpointy *-auto (lub flag ?auto=1 do istniejących).",
        "components/Gen4*.tsx — UI komponenty dostają dwie pary przycisków."
    ])

    H2(doc, "8.2 Sekwencja wywołań (FR-2, edycja strony)")
    CODE(doc,
"""User → click "✨ Zastosuj"
   ↓
[Browser] POST /api/v4/pages/[pageId]/ai-edit { instruction }
   ↓
[Next.js Route Handler]
   ├─ authenticate(req) → email
   ├─ ownPage(pageId, email) → bool
   ├─ buildPageEditPrompt(pageId, instruction) → { system, user, ... }
   ├─ callClaude({ system, user, model: EDIT_MODEL, maxTokens: 16000 })
   │      └─ Anthropic API: POST /v1/messages
   ├─ parsePageEditResponse(ai.text) → ParsedElements
   ├─ replacePageElements(pageId, parsed) → count
   └─ INSERT gen4_ai_history (workflow_type='page_edit', cost_usd, ...)
   ↓
[Browser] reload elements → re-render canvas""")

    H2(doc, "8.3 Decyzje techniczne")
    BULLETS(doc, [
        "Synchroniczne wywołanie API (nie streaming) — UX dla operacji < 60 s "
        "jest wystarczająco dobry z prostym loaderem.",
        "Brak per-user kolejki — Anthropic tier 4 daje 3000 RPM, wystarczy.",
        "Brak własnej kolejki retry w bazie — retry inline w Node, max 3 próby. "
        "Jeśli padnie, zapisujemy w ai_log i pokazujemy błąd. Brak długich retry "
        "ogonów (zatrzymują UX).",
        "Brak SSE / Webhooks — wszystko sync. Jeśli ktoś zamknie kartę, request "
        "może i tak dobiec do końca (Vercel kończy maxDuration=120s)."
    ])

    # ─── 9. Plan etapów ────────────────────────────────────────────────────
    H1(doc, "9. Plan etapów implementacji")
    add_table_with_header(doc,
        headers=["Etap", "Zakres", "Oszacowanie", "Kryterium ukończenia"],
        rows=[
            ["1. Klucz + telemetria", "Dodanie ANTHROPIC_API_KEY do Vercel; rozszerzenie callClaude o cost calculation + retry; rozbudowa zapisu gen4_ai_history (workflow_type, cost_usd).", "0,5 dnia", "Wywołanie z dev → wpis w gen4_ai_history z kosztem."],
            ["2. Auto-generacja (FR-1)", "Aktywacja istniejącego bloku 'if (apiKey)' w generate route. Pełen test end-to-end (wizard → ready w < 60 s).", "0,5 dnia", "Nowy projekt generowany 1 kliknięciem, ~10 stron."],
            ["3. Auto-edycja strony (FR-2)", "Nowy endpoint /pages/[pageId]/ai-edit. UI dual-button w Gen4Editor.tsx (PageAiAssistant).", "1 dzień", "Klik 'Zastosuj' → strona zaktualizowana w < 15 s."],
            ["4. Auto Apply DS (FR-3)", "Endpointy /projects/[id]/apply-design i /pages/[pageId]/apply-design. UI w Gen4DesignSystemPanel.", "1 dzień", "Klik 'Zastosuj DS do projektu' → wszystkie strony przepisane w < 90 s."],
            ["5. Auto tłumaczenia (FR-4)", "Endpoint /projects/[id]/translate. UI w Gen4TranslationsPanel.", "1 dzień", "Wybór języka + 'Generuj' → translations table zapełniona w < 30 s."],
            ["6. Dashboard kosztowy (FR-6)", "Sekcja Project Info pokazuje koszty + ostatnie wywołania.", "0,5 dnia", "Suma USD widoczna w UI; zgodna z ai_history."],
            ["7. Rate limiting + alerty (NFR-7.1, 7.5)", "Sliding-window rate limit per user; cron alert dzienny.", "0,5 dnia", "60+ wywołań w godzinie → HTTP 429; alert e-mail przy > 10 USD/dzień."],
            ["8. Hardening + dokumentacja", "Testy E2E (Playwright lub manual); aktualizacja README z workflow.", "0,5 dnia", "Wszystkie 5 workflow-ów udokumentowane; smoke test green."],
        ],
        col_widths_cm=[3, 7, 2, 4])
    P(doc, "Łącznie: ~5–6 dni roboczych. Przy intensywnej pracy: 1 sprint (2 tyg).")

    # ─── 10. Ryzyka ────────────────────────────────────────────────────────
    H1(doc, "10. Ryzyka i mitigacje")
    add_table_with_header(doc,
        headers=["Ryzyko", "Prawdopodobieństwo", "Wpływ", "Mitigacja"],
        rows=[
            ["Wyczerpanie klucza API / rate limit", "Średnie", "Wysoki", "Fallback do manual; alert na próg zużycia (75% miesięcznego budżetu)."],
            ["AI generuje halucynacje (wymyślone numery SAR, NIP)", "Niskie (dzięki zasadzie DO UZUPEŁNIENIA)", "Wysoki (compliance)", "Code review promptów; manualna weryfikacja sekcji prawnych przed drukiem; faza 2 — extraction z plików referencyjnych zastępuje placeholdery."],
            ["Niekontrolowany wzrost kosztów", "Średnie", "Średni", "Dashboard kosztowy + alert; cap dzienny (np. 5 USD); użycie Haiku 4.5 zamiast Sonnet 4.6 gdy możliwe."],
            ["Niska jakość JSON-a (błąd parsowania)", "Niskie", "Średni", "Już mamy escapeControlCharsInStrings; alternative: prompt z 'tool use' wymuszający schemę odpowiedzi."],
            ["Anthropic deprecation modelu", "Niskie (w horyzoncie 6 mies.)", "Niski", "INITIAL_MODEL / EDIT_MODEL jako stałe — pojedyncza zmiana."],
            ["Wyciek klucza API (PR z .env)", "Niskie", "Wysoki", ".env* w .gitignore (już jest); klucz tylko w Vercel Encrypted; rotacja co 6 mies."],
        ],
        col_widths_cm=[5, 3.5, 2.5, 6])

    # ─── 11. Decyzje otwarte ───────────────────────────────────────────────
    H1(doc, "11. Decyzje otwarte")
    BULLETS(doc, [
        "Czy używać prompt caching API (Anthropic) dla powtarzających się "
        "promptów systemowych? Decyzja po zebraniu telemetrii z fazy 1.",
        "Czy dorzucić generację obrazków (placeholdery → wyrenderowane ikony) "
        "w fazie 2? Wymaga decyzji o providerze (Imagen, DALL-E 3, Flux). "
        "Out-of-scope dla obecnego PRD.",
        "Czy budować retry workera (jobs w Supabase) zamiast inline retry? "
        "Tylko jeśli error rate przekroczy 5% — wtedy refactor.",
        "Czy ekspozycja \"Regeneruj losowo\" (różny seed/temperature)? "
        "Heurystycznie tak, ale potrzebny mały eksperyment by zobaczyć czy "
        "różnice są wartościowe."
    ])

    # ─── 12. Załączniki ────────────────────────────────────────────────────
    H1(doc, "12. Załączniki")
    H2(doc, "12.1 Mapowanie endpointów manual → auto")
    add_table_with_header(doc,
        headers=["Manual prompt endpoint", "Manual import endpoint", "Nowy auto endpoint", "Model"],
        rows=[
            ["GET /api/v4/projects/[id]/prompt", "POST /api/v4/projects/[id]/import", "POST /api/v4/projects/generate (już istnieje, aktywacja API path)", "Sonnet 4.6"],
            ["POST /api/v4/pages/[pageId]/edit-prompt", "POST /api/v4/pages/[pageId]/replace-elements", "POST /api/v4/pages/[pageId]/ai-edit (NOWE)", "Haiku 4.5"],
            ["POST /api/v4/projects/[id]/apply-design-prompt", "POST /api/v4/projects/[id]/import", "POST /api/v4/projects/[id]/apply-design (NOWE)", "Sonnet 4.6"],
            ["POST /api/v4/pages/[pageId]/apply-design-prompt", "POST /api/v4/pages/[pageId]/replace-elements", "POST /api/v4/pages/[pageId]/apply-design (NOWE)", "Haiku 4.5"],
            ["GET /api/v4/projects/[id]/translate-prompt", "POST /api/v4/projects/[id]/translations", "POST /api/v4/projects/[id]/translate (NOWE)", "Haiku 4.5"],
        ],
        col_widths_cm=[5, 5, 5, 2])

    H2(doc, "12.2 Zmienne środowiskowe wymagane na Vercel")
    add_table_with_header(doc,
        headers=["Zmienna", "Wartość", "Notatka"],
        rows=[
            ["ANTHROPIC_API_KEY", "sk-ant-...", "Encrypted, Production + Preview"],
            ["SUPABASE_URL", "https://....supabase.co", "Już skonfigurowany"],
            ["SUPABASE_SERVICE_ROLE_KEY", "eyJ...", "Już skonfigurowany"],
            ["INTERNAL_PROXY_SECRET", "...", "Już skonfigurowany"],
            ["AUTH_DOMAIN", "locon.pl", "Już skonfigurowany"],
            ["JWT_SECRET", "...", "Już skonfigurowany"],
        ],
        col_widths_cm=[5, 6, 6])

    H2(doc, "12.3 Glosariusz")
    BULLETS(doc, [
        "DS — Design System (gen4_design_systems).",
        "FR — Functional Requirement.",
        "NFR — Non-Functional Requirement.",
        "RPM — Requests per Minute (limit API).",
        "TPM — Tokens per Minute (limit API).",
        "SSE — Server-Sent Events (sposób streamingu)."
    ])

    # ─── Stopka ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"PRD wygenerowany {date.today().isoformat()} — wersja 1.0")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.save(OUTPUT_PATH)
    print(f"PRD zapisany: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
